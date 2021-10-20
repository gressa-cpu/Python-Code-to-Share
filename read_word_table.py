#Reads word document table info 
#Made by: Gressa
#Last Modified : 22/08/2021


def main():
    
    import docx
    import pandas as pd
    from docx.api import Document
    

    #Load the document

    #doc = Document('example03.docx')
    #Trying another document
    doc = Document('PandasTableExtraction.docx')

    table = doc.tables[0]

    print(table)


    #Store data
    data = []

    keys = None

    def table_with_header(table, data):
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            
            #Map the data based on first row
            #Headers will become keys to your dictionary
            if i == 0:
                keys = tuple(text)
                continue
            
            #Construct a dict for this row, mapping keys to values for this row
            row_data = dict(zip(keys, text))
            data.append(row_data)
        return data

    #If you want tuple for each row




    def row_to_tuple(table, data):                          #Return a list of tuples of every row in table. 
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            #Map the data based on first row
            #Headers will become keys to your dictionary
            
            #You can uncomment the following if statement if you'd like to use the first row as header. 
            #Be careful cause then it wouldn't return the first row in the tuple. 
            # if i == 0:
            #     keys = tuple(text)
            #     continue

            # Construct a tuple for this row
            row_data = tuple(text)
            data.append(row_data)
        return data
            

    
    def read_using_panda(doc, table_num = 1, num_header = 1 ):
        table = doc.tables[table_num - 1]
        data = [[cell.text for cell in row.cells] for row in table.rows]
        #Panda Data Frame
        df = pd.DataFrame(data)
        
        #Initial data frame should be formatted according to the scenario in hand
        
        #For example with a single header scenario
        if num_header ==1:
            df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop = True)
        
        #For two headers scenario
        elif num_header == 2:
            outside_col, inside_col = df.iloc[0], df.iloc[1]
            hier_index = pd.MultiIndex.from_tuples(list(zip(outside_col, inside_col)))
            df = pd.DataFrame(data, columns= hier_index).drop(df.index[[0, 1]]).reset_index(drop = True)
            
        elif num_header > 2:
            print("More than two headers not currently supported")
            df = pd.DataFrame()
        
        
        return df
    
        
    def convert_panda_df_to_excel(data, name = 'output.xlsx'):
        #This function takes a pandas dataframe and saves it as an excel file. 
        data.to_excel(name)   
    
    
    
    #row_to_tuple(table, data)

    data = read_using_panda(doc, 3 , 2)
    print(data)
    
    convert_panda_df_to_excel(data, 'book.xlsx')
    
    
if __name__ == '__main__':
    main()
